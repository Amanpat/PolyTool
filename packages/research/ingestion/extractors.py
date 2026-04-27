"""RIS v1 ingestion — extractor ABC and PlainTextExtractor.

Extractors convert a source (file path or raw text) into an ExtractedDocument
dataclass ready for hard-stop checking and pipeline ingestion.
"""

from __future__ import annotations

import hashlib
import re
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from packages.research.evaluation.types import SOURCE_FAMILIES

# ---------------------------------------------------------------------------
# Optional dependency detection (module-level, non-raising)
# ---------------------------------------------------------------------------

try:
    import pdfplumber as _pdfplumber
except ImportError:
    _pdfplumber = None  # type: ignore[assignment]

try:
    import docx as _docx
except ImportError:
    _docx = None  # type: ignore[assignment]


# ---------------------------------------------------------------------------
# Data types
# ---------------------------------------------------------------------------

@dataclass
class ExtractedDocument:
    """Intermediate representation produced by an extractor.

    Attributes
    ----------
    title:
        Human-readable document title.
    body:
        Full document body text.
    source_url:
        Origin URL or internal reference (e.g. ``file:///abs/path``).
    source_family:
        Source-family key matching freshness_decay.json entries.
    author:
        Document author (default: "unknown").
    publish_date:
        ISO-8601 string of the publication date, or None if unknown.
    metadata:
        Arbitrary extra key/value pairs (e.g. ``content_hash``).
    """
    title: str
    body: str
    source_url: str
    source_family: str
    author: str = "unknown"
    publish_date: Optional[str] = None
    metadata: dict = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Extractor ABC
# ---------------------------------------------------------------------------

class Extractor(ABC):
    """Abstract base class for document extractors."""

    @abstractmethod
    def extract(self, source: "str | Path", **kwargs) -> ExtractedDocument:
        """Extract a document from *source*.

        Parameters
        ----------
        source:
            A file path (``str`` or ``Path``) or raw text string.
        **kwargs:
            Extractor-specific options (e.g. ``title``, ``source_type``,
            ``author``, ``publish_date``).

        Returns
        -------
        ExtractedDocument
        """
        ...


# ---------------------------------------------------------------------------
# PlainTextExtractor
# ---------------------------------------------------------------------------

_H1_RE = re.compile(r"^#\s+(.+)", re.MULTILINE)


def _sha256_hex(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


class PlainTextExtractor(Extractor):
    """Extract plain text and Markdown files (or raw strings).

    File mode (default)
    -------------------
    Pass a ``Path`` or a string that points to an existing file.
    - Title: first Markdown H1 (``# Title``) if present; else filename stem.
    - Body: entire file content (utf-8).
    - source_url: ``file://{absolute_path}``.

    Raw-text mode
    -------------
    Pass a string that does NOT point to an existing file.
    - ``title`` kwarg is **required**.
    - Body: the string itself.
    - source_url: ``internal://manual``.

    Common kwargs
    -------------
    source_type : str
        Source type key (default: ``"manual"``).  Mapped to ``source_family``
        via ``SOURCE_FAMILIES``; falls back to the source_type itself.
    author : str
        Document author (default: ``"unknown"``).
    publish_date : str or None
        ISO-8601 publication date (default: ``None``).
    title : str
        Required in raw-text mode; optional override in file mode.
    """

    def extract(self, source: "str | Path", **kwargs) -> ExtractedDocument:  # type: ignore[override]
        source_type: str = kwargs.get("source_type", "manual")
        author: str = kwargs.get("author", "unknown")
        publish_date: Optional[str] = kwargs.get("publish_date", None)

        # Determine source_family from source_type via mapping
        source_family: str = SOURCE_FAMILIES.get(source_type, source_type)

        # --- decide mode ---
        as_path = Path(source) if isinstance(source, str) else source
        if as_path.exists():
            # File mode
            body = as_path.read_text(encoding="utf-8")
            abs_path = str(as_path.resolve()).replace("\\", "/")
            source_url = f"file://{abs_path}"

            # Title from H1 or filename stem
            title_override = kwargs.get("title")
            if title_override:
                title = title_override
            else:
                m = _H1_RE.search(body)
                title = m.group(1).strip() if m else as_path.stem

        else:
            # Raw-text mode: source must be a non-file string
            if isinstance(source, Path):
                raise FileNotFoundError(f"No such file: {source}")
            # Check if caller meant a file path that doesn't exist
            if isinstance(source, str) and "/" in source or (isinstance(source, str) and "\\" in source):
                # Looks like a path that doesn't exist
                raise FileNotFoundError(f"No such file: {source}")

            title = kwargs.get("title")
            if not title:
                raise ValueError(
                    "PlainTextExtractor: 'title' kwarg is required in raw-text mode."
                )
            body = str(source)
            source_url = "internal://manual"

        content_hash = _sha256_hex(body)
        metadata = {"content_hash": content_hash}

        return ExtractedDocument(
            title=title,
            body=body,
            source_url=source_url,
            source_family=source_family,
            author=author,
            publish_date=publish_date,
            metadata=metadata,
        )


# ---------------------------------------------------------------------------
# MarkdownExtractor
# ---------------------------------------------------------------------------


class MarkdownExtractor(Extractor):
    """Extract Markdown files into ExtractedDocument.

    Delegates entirely to PlainTextExtractor — Markdown files are plain text
    and the H1-title extraction already works correctly there.  This class
    exists so callers can request ``get_extractor('markdown')`` and get an
    explicitly named extractor rather than the generic PlainTextExtractor.
    """

    def __init__(self) -> None:
        self._delegate = PlainTextExtractor()

    def extract(self, source: "str | Path", **kwargs) -> ExtractedDocument:  # type: ignore[override]
        return self._delegate.extract(source, **kwargs)


# ---------------------------------------------------------------------------
# StructuredMarkdownExtractor
# ---------------------------------------------------------------------------

# Heading pattern: matches H1-H6 headings (e.g. # H1, ## H2, ### H3 ...)
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)", re.MULTILINE)
# Table row pattern: line is a table row (starts/ends with pipe)
_TABLE_ROW_RE = re.compile(r"^\|.+\|", re.MULTILINE)
# Table separator row: `|---|---|` or `|:---:|`
_TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|", re.MULTILINE)
# Fenced code block: ``` or ~~~
_FENCE_RE = re.compile(r"^(`{3,}|~{3,})", re.MULTILINE)


def _count_tables(text: str) -> int:
    """Count Markdown table blocks in *text*.

    A table block is defined as a sequence of pipe-delimited rows that includes
    at least one separator row (``|---|---|``).  Adjacent table rows are counted
    as part of the same table.
    """
    # A table has >= 1 separator rows; count distinct table blocks by
    # detecting transitions from non-table to table content.
    lines = text.splitlines()
    in_table = False
    has_sep = False
    table_count = 0
    for line in lines:
        stripped = line.strip()
        is_row = bool(re.match(r"^\|.+\|", stripped))
        is_sep = bool(re.match(r"^\|[\s\-:|]+\|", stripped))

        if is_row:
            if not in_table:
                in_table = True
                has_sep = False
            if is_sep:
                has_sep = True
        else:
            if in_table and has_sep:
                table_count += 1
            in_table = False
            has_sep = False

    # Handle table at end of file
    if in_table and has_sep:
        table_count += 1

    return table_count


def _count_code_blocks(text: str) -> int:
    """Count fenced code blocks (``` or ~~~) in *text*."""
    fence_matches = _FENCE_RE.findall(text)
    # Each block is an open + close fence pair
    return len(fence_matches) // 2


class StructuredMarkdownExtractor(Extractor):
    """Structure-aware Markdown extractor.

    Extends PlainTextExtractor by parsing Markdown structural elements and
    storing them as rich metadata:

    - ``sections``: list of section heading title strings
    - ``section_count``: number of headings found
    - ``header_count``: total heading lines (same as section_count, both kept
      for clarity — header_count counts raw heading lines including any that
      may repeat)
    - ``table_count``: number of Markdown table blocks detected
    - ``code_block_count``: number of fenced code blocks detected

    The body text is returned UNCHANGED (structure is preserved, not stripped).
    The value-add is the metadata.

    Falls back to PlainTextExtractor behaviour on any parsing exception.
    """

    def __init__(self) -> None:
        self._delegate = PlainTextExtractor()

    def extract(self, source: "str | Path", **kwargs) -> ExtractedDocument:  # type: ignore[override]
        # First, get a base document from PlainTextExtractor
        doc = self._delegate.extract(source, **kwargs)

        # Attempt structural parsing of body
        try:
            body = doc.body
            heading_matches = _HEADING_RE.findall(body)
            # heading_matches: list of (hashes, title) tuples
            sections = [title.strip() for _, title in heading_matches]
            section_count = len(sections)
            header_count = section_count  # all headings are headers

            table_count = _count_tables(body)
            code_block_count = _count_code_blocks(body)

            structural_metadata = {
                "sections": sections,
                "section_count": section_count,
                "header_count": header_count,
                "table_count": table_count,
                "code_block_count": code_block_count,
            }
            doc.metadata.update(structural_metadata)
        except Exception:
            # Graceful fallback: return doc with base metadata only (no crash)
            pass

        return doc


# ---------------------------------------------------------------------------
# PDFExtractor — real implementation with graceful ImportError fallback
# ---------------------------------------------------------------------------


class PDFExtractor(Extractor):
    """PDF extractor using pdfplumber (optional dependency).

    If ``pdfplumber`` is installed, extracts text from all pages and computes
    structural metadata.  If pdfplumber is NOT installed, raises ``ImportError``
    with a helpful install hint rather than ``NotImplementedError``.

    Attributes
    ----------
    _pdfplumber:
        The pdfplumber module (set at init time from module-level import).
        Can be set to ``None`` in tests to simulate the missing-dependency path.
    """

    def __init__(self) -> None:
        self._pdfplumber = _pdfplumber

    def extract(self, source: "str | Path", **kwargs) -> ExtractedDocument:  # type: ignore[override]
        if self._pdfplumber is None:
            raise ImportError(
                "PDF extraction requires pdfplumber. "
                "Install: pip install pdfplumber"
            )

        source_path = Path(source)
        if not source_path.exists():
            raise FileNotFoundError(f"No such file: {source_path}")

        source_type: str = kwargs.get("source_type", "manual")
        author: str = kwargs.get("author", "unknown")
        publish_date: Optional[str] = kwargs.get("publish_date", None)
        source_family: str = SOURCE_FAMILIES.get(source_type, source_type)

        with self._pdfplumber.open(str(source_path)) as pdf:
            pages = pdf.pages
            page_count = len(pages)
            texts = []
            for page in pages:
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text)

        body = "\n\n".join(texts)

        # Title: from first non-empty line of first page, or filename stem
        title_override = kwargs.get("title")
        if title_override:
            title = title_override
        else:
            first_line = texts[0].strip().splitlines()[0].strip() if texts else ""
            title = first_line if first_line else source_path.stem

        abs_path = str(source_path.resolve()).replace("\\", "/")
        source_url = f"file://{abs_path}"
        content_hash = _sha256_hex(body)
        metadata = {
            "content_hash": content_hash,
            "page_count": page_count,
        }

        return ExtractedDocument(
            title=title,
            body=body,
            source_url=source_url,
            source_family=source_family,
            author=author,
            publish_date=publish_date,
            metadata=metadata,
        )


# ---------------------------------------------------------------------------
# MarkerPDFExtractor — structured Markdown via marker-pdf (optional dep)
# ---------------------------------------------------------------------------

_MARKER_METADATA_SIZE_LIMIT = 20 * 1024 * 1024  # 20 MB


class MarkerPDFExtractor(Extractor):
    """PDF extractor using marker-pdf for structured Markdown output.

    marker-pdf is an optional heavy dependency (pulls PyTorch/model weights).
    This extractor imports marker lazily so it never affects startup time.

    Install: ``pip install 'polytool[ris-marker]'``

    Parameters
    ----------
    _marker_modules:
        Injectable dict with keys ``"PdfConverter"``, ``"create_model_dict"``,
        ``"text_from_rendered"``.  Used by offline tests to bypass real imports.
    _enable_llm:
        Reserved LLM-boost flag.  When ``True``, sets ``body_source``
        to ``"marker_llm_boost"``.  Default ``False`` — no external LLM calls
        are made regardless.  Can also be activated via ``RIS_MARKER_LLM=1``.
    """

    def __init__(
        self,
        _marker_modules: "Optional[dict]" = None,
        _enable_llm: bool = False,
    ) -> None:
        self._marker_modules = _marker_modules
        import os as _os
        self._enable_llm = _enable_llm or _os.environ.get("RIS_MARKER_LLM", "").lower() in ("1", "true")

    def _load_marker(self) -> dict:
        """Lazily import marker-pdf modules. Returns dict of callables."""
        if self._marker_modules is not None:
            return self._marker_modules
        try:
            from marker.converters.pdf import PdfConverter
            from marker.models import create_model_dict
            from marker.output import text_from_rendered
            return {
                "PdfConverter": PdfConverter,
                "create_model_dict": create_model_dict,
                "text_from_rendered": text_from_rendered,
            }
        except ImportError as exc:
            raise ImportError(
                "Marker PDF extraction requires marker-pdf. "
                "Install: pip install 'polytool[ris-marker]'"
            ) from exc

    def _discover_version(self) -> "Optional[str]":
        try:
            import importlib.metadata as _imeta
            return _imeta.version("marker-pdf")
        except Exception:
            return None

    def extract(self, source: "str | Path", **kwargs) -> ExtractedDocument:  # type: ignore[override]
        import json as _json

        mods = self._load_marker()  # raises ImportError if marker not installed

        source_path = Path(source)
        if not source_path.exists():
            raise FileNotFoundError(f"No such file: {source_path}")

        source_type: str = kwargs.get("source_type", "manual")
        author: str = kwargs.get("author", "unknown")
        publish_date: "Optional[str]" = kwargs.get("publish_date", None)
        source_family: str = SOURCE_FAMILIES.get(source_type, source_type)

        marker_version = self._discover_version()

        PdfConverter = mods["PdfConverter"]
        create_model_dict = mods["create_model_dict"]
        text_from_rendered = mods["text_from_rendered"]

        model_dict = create_model_dict()
        converter = PdfConverter(artifact_dict=model_dict)
        rendered = converter(str(source_path))

        result = text_from_rendered(rendered)
        if isinstance(result, tuple) and result:
            markdown_text = result[0] or ""
            # Second element is out_meta if it's a plain dict (not image dict)
            out_meta: dict = {}
            for item in result[1:]:
                if isinstance(item, dict) and not any(
                    isinstance(v, (bytes, bytearray)) for v in item.values()
                ):
                    out_meta = item
                    break
        else:
            markdown_text = str(result) if result else ""
            out_meta = {}

        # Strip image binaries from out_meta
        structured_metadata = {
            k: v for k, v in out_meta.items()
            if k != "images" and not isinstance(v, (bytes, bytearray))
        }

        # Enforce 20 MB size cap on structured_metadata
        structured_metadata_truncated = False
        try:
            if len(_json.dumps(structured_metadata).encode("utf-8")) > _MARKER_METADATA_SIZE_LIMIT:
                structured_metadata_truncated = True
                structured_metadata = {
                    "truncated": True,
                    "page_count": structured_metadata.get("page_count", 0),
                    "truncation_reason": f"metadata exceeded {_MARKER_METADATA_SIZE_LIMIT // (1024 * 1024)}MB cap",
                }
        except (TypeError, ValueError):
            structured_metadata = {}

        # Extract page count
        page_count: int = int(out_meta.get("page_count", 0) or 0)
        if not page_count:
            try:
                page_count = int(rendered.metadata.page_count)
            except (AttributeError, TypeError, ValueError):
                pass

        title_override = kwargs.get("title")
        if title_override:
            title = str(title_override)
        else:
            m = _H1_RE.search(markdown_text)
            title = m.group(1).strip() if m else source_path.stem

        abs_path = str(source_path.resolve()).replace("\\", "/")
        source_url = f"file://{abs_path}"
        content_hash = _sha256_hex(markdown_text)

        # LLM-boost is not yet wired: always report body_source="marker".
        # If _enable_llm was requested, record that as metadata signals so
        # operators can see the intent without a misleading body_source label.
        import logging as _logging
        _log = _logging.getLogger(__name__)
        if self._enable_llm:
            _log.warning(
                "RIS_MARKER_LLM / _enable_llm=True is set but no LLM call is "
                "wired in MarkerPDFExtractor. Reporting body_source='marker' "
                "with marker_llm_requested=True, marker_llm_applied=False."
            )

        metadata_out: dict = {
            "content_hash": content_hash,
            "page_count": page_count,
            "body_source": "marker",
            "has_structured_metadata": True,
            "structured_metadata": structured_metadata,
            "structured_metadata_truncated": structured_metadata_truncated,
        }
        if self._enable_llm:
            metadata_out["marker_llm_requested"] = True
            metadata_out["marker_llm_applied"] = False
        if marker_version is not None:
            metadata_out["marker_version"] = marker_version

        return ExtractedDocument(
            title=title,
            body=markdown_text,
            source_url=source_url,
            source_family=source_family,
            author=author,
            publish_date=publish_date,
            metadata=metadata_out,
        )


# ---------------------------------------------------------------------------
# DocxExtractor — real implementation with graceful ImportError fallback
# ---------------------------------------------------------------------------


class DocxExtractor(Extractor):
    """DOCX extractor using python-docx (optional dependency).

    If ``python-docx`` is installed, extracts paragraph text and computes
    basic metadata.  If python-docx is NOT installed, raises ``ImportError``
    with a helpful install hint rather than ``NotImplementedError``.

    Attributes
    ----------
    _docx:
        The docx module (set at init time from module-level import).
        Can be set to ``None`` in tests to simulate the missing-dependency path.
    """

    def __init__(self) -> None:
        self._docx = _docx

    def extract(self, source: "str | Path", **kwargs) -> ExtractedDocument:  # type: ignore[override]
        if self._docx is None:
            raise ImportError(
                "DOCX extraction requires python-docx. "
                "Install: pip install python-docx"
            )

        source_path = Path(source)
        if not source_path.exists():
            raise FileNotFoundError(f"No such file: {source_path}")

        source_type: str = kwargs.get("source_type", "manual")
        author: str = kwargs.get("author", "unknown")
        publish_date: Optional[str] = kwargs.get("publish_date", None)
        source_family: str = SOURCE_FAMILIES.get(source_type, source_type)

        document = self._docx.Document(str(source_path))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        paragraph_count = len(paragraphs)
        body = "\n\n".join(paragraphs)

        # Title: first heading paragraph or filename stem
        title_override = kwargs.get("title")
        if title_override:
            title = title_override
        else:
            title = source_path.stem
            for p in document.paragraphs:
                if p.style.name.startswith("Heading") and p.text.strip():
                    title = p.text.strip()
                    break

        abs_path = str(source_path.resolve()).replace("\\", "/")
        source_url = f"file://{abs_path}"
        content_hash = _sha256_hex(body)
        metadata = {
            "content_hash": content_hash,
            "paragraph_count": paragraph_count,
        }

        return ExtractedDocument(
            title=title,
            body=body,
            source_url=source_url,
            source_family=source_family,
            author=author,
            publish_date=publish_date,
            metadata=metadata,
        )


# ---------------------------------------------------------------------------
# Stub extractors (kept for backward compatibility; tests may still reference them)
# ---------------------------------------------------------------------------


class StubPDFExtractor(Extractor):
    """Stub PDF extractor — raises NotImplementedError.

    Kept for backward compatibility. Use PDFExtractor for real extraction.
    """

    def extract(self, source: "str | Path", **kwargs) -> ExtractedDocument:  # type: ignore[override]
        raise NotImplementedError(
            "PDF extraction requires an external library. "
            "Install and configure one of: docling, marker, pymupdf4llm. "
            "Replace StubPDFExtractor with a real implementation once chosen."
        )


class StubDocxExtractor(Extractor):
    """Stub DOCX extractor — raises NotImplementedError.

    Kept for backward compatibility. Use DocxExtractor for real extraction.
    """

    def extract(self, source: "str | Path", **kwargs) -> ExtractedDocument:  # type: ignore[override]
        raise NotImplementedError(
            "DOCX extraction requires python-docx. "
            "Install it (pip install python-docx) and replace StubDocxExtractor "
            "with a real implementation."
        )


# ---------------------------------------------------------------------------
# Extractor registry and factory
# ---------------------------------------------------------------------------

EXTRACTOR_REGISTRY: dict[str, type[Extractor]] = {
    "plain_text": PlainTextExtractor,
    "markdown": MarkdownExtractor,
    "structured_markdown": StructuredMarkdownExtractor,
    "pdf": PDFExtractor,
    "marker_pdf": MarkerPDFExtractor,
    "docx": DocxExtractor,
}


def get_extractor(name: str) -> Extractor:
    """Return an instantiated extractor for *name*.

    Parameters
    ----------
    name:
        Key in ``EXTRACTOR_REGISTRY`` (e.g. ``"plain_text"``, ``"markdown"``).

    Returns
    -------
    Extractor
        A new instance of the requested extractor class.

    Raises
    ------
    KeyError
        If *name* is not in ``EXTRACTOR_REGISTRY``.
    """
    cls = EXTRACTOR_REGISTRY[name]
    return cls()
